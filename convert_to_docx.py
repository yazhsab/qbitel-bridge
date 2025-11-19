#!/usr/bin/env python3
"""
Professional DOCX Converter for QBITEL AI Investor Documents
Converts Markdown to editable Microsoft Word documents with enterprise styling
"""

import os
import sys
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from html.parser import HTMLParser


class MarkdownToDOCXConverter:
    """Convert Markdown to professionally formatted DOCX"""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure professional document styles"""
        # Configure Normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Heading 1
        style = self.doc.styles['Heading 1']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(24)
        font.bold = True
        font.color.rgb = RGBColor(0, 102, 204)  # QBITEL Blue

        # Heading 2
        style = self.doc.styles['Heading 2']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(18)
        font.bold = True
        font.color.rgb = RGBColor(0, 102, 204)

        # Heading 3
        style = self.doc.styles['Heading 3']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(14)
        font.bold = True
        font.color.rgb = RGBColor(44, 62, 80)

    def _add_cover_page(self, title):
        """Add professional cover page"""
        # Company name
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('QBITEL AI')
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)

        self.doc.add_paragraph()  # Spacing

        # Document title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(44, 62, 80)

        self.doc.add_paragraph()  # Spacing

        # Confidential notice
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('CONFIDENTIAL & PROPRIETARY\nFor Investor Use Only')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(192, 0, 0)
        run.font.bold = True

        self.doc.add_paragraph()  # Spacing

        # Date
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(datetime.now().strftime('%B %Y'))
        run.font.size = Pt(12)

        # Page break
        self.doc.add_page_break()

    def _add_footer(self):
        """Add footer with page numbers"""
        section = self.doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = "© 2025 QBITEL AI. All rights reserved. | Confidential"

    def convert_html_to_docx(self, html_content):
        """Convert HTML (from markdown) to DOCX elements"""
        from html.parser import HTMLParser

        class DocxHTMLParser(HTMLParser):
            def __init__(self, doc):
                super().__init__()
                self.doc = doc
                self.current_paragraph = None
                self.in_table = False
                self.table_data = []
                self.current_row = []
                self.list_level = 0

            def handle_starttag(self, tag, attrs):
                if tag == 'h1':
                    self.current_paragraph = self.doc.add_heading(level=1)
                elif tag == 'h2':
                    self.current_paragraph = self.doc.add_heading(level=2)
                elif tag == 'h3':
                    self.current_paragraph = self.doc.add_heading(level=3)
                elif tag == 'h4':
                    self.current_paragraph = self.doc.add_heading(level=4)
                elif tag == 'p':
                    self.current_paragraph = self.doc.add_paragraph()
                elif tag == 'ul' or tag == 'ol':
                    self.list_level += 1
                elif tag == 'li':
                    self.current_paragraph = self.doc.add_paragraph(style='List Bullet' if self.list_level > 0 else 'Normal')
                elif tag == 'table':
                    self.in_table = True
                    self.table_data = []
                elif tag == 'tr':
                    self.current_row = []
                elif tag == 'th' or tag == 'td':
                    pass  # Handle in handle_data
                elif tag == 'strong' or tag == 'b':
                    pass  # Handle in handle_data with bold
                elif tag == 'code':
                    pass  # Handle in handle_data with code style
                elif tag == 'hr':
                    self.doc.add_paragraph('_' * 80)

            def handle_endtag(self, tag):
                if tag in ['h1', 'h2', 'h3', 'h4', 'p', 'li']:
                    self.current_paragraph = None
                elif tag == 'ul' or tag == 'ol':
                    self.list_level -= 1
                elif tag == 'tr':
                    if self.current_row:
                        self.table_data.append(self.current_row)
                elif tag == 'table':
                    if self.table_data:
                        self._create_table()
                    self.in_table = False

            def handle_data(self, data):
                data = data.strip()
                if not data:
                    return

                if self.in_table:
                    self.current_row.append(data)
                elif self.current_paragraph:
                    self.current_paragraph.add_run(data)
                else:
                    # Fallback: create new paragraph
                    p = self.doc.add_paragraph()
                    p.add_run(data)

            def _create_table(self):
                if not self.table_data:
                    return

                num_cols = max(len(row) for row in self.table_data)
                table = self.doc.add_table(rows=len(self.table_data), cols=num_cols)
                table.style = 'Light Grid Accent 1'

                for i, row_data in enumerate(self.table_data):
                    row = table.rows[i]
                    for j, cell_data in enumerate(row_data):
                        if j < len(row.cells):
                            cell = row.cells[j]
                            cell.text = cell_data
                            # Make header row bold
                            if i == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                                        run.font.color.rgb = RGBColor(255, 255, 255)
                                cell._element.get_or_add_tcPr().append(
                                    cell._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                                )

                self.doc.add_paragraph()  # Spacing after table
                self.table_data = []

        parser = DocxHTMLParser(self.doc)
        parser.feed(html_content)


def convert_markdown_to_docx(markdown_file: Path, output_dir: Path = None):
    """
    Convert a markdown file to a professionally formatted DOCX

    Args:
        markdown_file: Path to the markdown file
        output_dir: Output directory for the DOCX
    """
    if output_dir is None:
        output_dir = markdown_file.parent

    output_dir.mkdir(parents=True, exist_ok=True)

    # Read markdown content
    print(f"Reading {markdown_file.name}...")
    with open(markdown_file, 'r', encoding='utf-8') as f:
        markdown_content = f.read()

    # Convert markdown to HTML first
    print(f"Converting {markdown_file.name} to HTML...")
    html_content = markdown.markdown(
        markdown_content,
        extensions=[
            'tables',
            'fenced_code',
            'nl2br',
            'sane_lists'
        ]
    )

    # Get document title
    title_map = {
        "INVESTOR_PITCH_DECK": "Series A Investment Deck",
        "PRODUCT_WALKTHROUGH": "Product & Technology Overview",
        "IMPLEMENTATION_PLAN_ROADMAP": "5-Year Implementation Roadmap"
    }
    title = title_map.get(markdown_file.stem, markdown_file.stem)

    # Create DOCX document
    print(f"Creating professional DOCX: {markdown_file.stem}.docx...")
    converter = MarkdownToDOCXConverter()

    # Add cover page
    converter._add_cover_page(title)

    # Convert HTML to DOCX
    converter.convert_html_to_docx(html_content)

    # Add footer
    converter._add_footer()

    # Save DOCX
    docx_filename = output_dir / f"{markdown_file.stem}.docx"
    converter.doc.save(docx_filename)

    print(f"[OK] Successfully created: {docx_filename}")
    print(f"   > Editable Microsoft Word document")
    print(f"   > Ready to share with investor teams")
    return docx_filename


def main():
    """Main function to convert all investor documents"""

    # Get the project root directory
    project_root = Path(__file__).parent

    # Files to convert
    files_to_convert = [
        "INVESTOR_PITCH_DECK.md",
        "PRODUCT_WALKTHROUGH.md",
        "IMPLEMENTATION_PLAN_ROADMAP.md"
    ]

    print("=" * 80)
    print("QBITEL AI - Enterprise DOCX Generator")
    print("Creating Editable Microsoft Word Documents")
    print("=" * 80)
    print(f"Project Root: {project_root}")
    print(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    print("=" * 80)
    print()

    output_dir = project_root / "investor_documents_docx"

    generated_files = []
    for filename in files_to_convert:
        markdown_file = project_root / filename
        if markdown_file.exists():
            try:
                docx_file = convert_markdown_to_docx(markdown_file, output_dir)
                generated_files.append(docx_file)
                print()
            except Exception as e:
                print(f"[ERROR] Error converting {filename}: {e}")
                import traceback
                traceback.print_exc()
                print()
        else:
            print(f"[WARNING] File not found: {filename}")
            print()

    print("=" * 80)
    print("Summary")
    print("=" * 80)
    print(f"Total files converted: {len(generated_files)}/{len(files_to_convert)}")
    print(f"Output directory: {output_dir}")
    print()
    print("Generated DOCX files:")
    for docx_file in generated_files:
        file_size = docx_file.stat().st_size / 1024  # KB
        print(f"  - {docx_file.name} ({file_size:.1f} KB)")
    print()
    print("BENEFITS:")
    print("  - Fully editable in Microsoft Word")
    print("  - Easy to share with investor teams")
    print("  - Can be commented on and collaborated")
    print("  - Professional enterprise formatting")
    print("  - Compatible with Google Docs")
    print("=" * 80)

    return len(generated_files) == len(files_to_convert)


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
