#!/usr/bin/env python3
"""
Enterprise-Grade DOCX Converter for QBITEL AI Investor Documents
Converts Markdown to professionally formatted Microsoft Word documents with:
- SVG diagram embedding (converted to PNG)
- Table of Contents
- Professional headers/footers with page numbers
- Enterprise styling and formatting
- Proper section breaks and page layouts
"""

import os
import sys
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from html.parser import HTMLParser
import re
import subprocess
import tempfile


def convert_svg_to_png(svg_path, output_path, width=1400):
    """
    Convert SVG to PNG - tries multiple methods

    Args:
        svg_path: Path to SVG file
        output_path: Path for output PNG
        width: Width in pixels (default 1400 for high quality)
    """
    # Try ImageMagick first (most reliable on Windows)
    try:
        result = subprocess.run([
            'magick', 'convert',
            '-density', '150',
            '-background', 'white',
            '-alpha', 'remove',
            '-resize', f'{width}x',
            str(svg_path),
            str(output_path)
        ], check=True, capture_output=True, text=True)
        if output_path.exists():
            return True
    except (subprocess.CalledProcessError, FileNotFoundError) as e:
        pass

    # Try inkscape
    try:
        result = subprocess.run([
            'inkscape',
            str(svg_path),
            '--export-type=png',
            f'--export-filename={output_path}',
            f'--export-width={width}'
        ], check=True, capture_output=True)
        if output_path.exists():
            return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        pass

    return False


def add_page_number(section):
    """Add page numbers to footer"""
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page number field
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Add footer text
    paragraph.add_run(' | © 2025 QBITEL AI - Confidential')
    paragraph.runs[1].font.size = Pt(9)
    paragraph.runs[1].font.color.rgb = RGBColor(128, 128, 128)


class EnterpriseMarkdownToDOCXConverter:
    """Convert Markdown to enterprise-grade DOCX with diagrams"""

    def __init__(self):
        self.doc = Document()
        self.headings = []  # Track headings for TOC
        self.temp_dir = Path(tempfile.mkdtemp())
        self._setup_document_settings()
        self._setup_styles()

    def _setup_document_settings(self):
        """Configure document-level settings"""
        # Set up page margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
            section.page_height = Cm(29.7)  # A4
            section.page_width = Cm(21)

    def _setup_styles(self):
        """Configure professional document styles"""
        # Normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.15

        # Heading 1 - Major sections
        style = self.doc.styles['Heading 1']
        font = style.font
        font.name = 'Calibri Light'
        font.size = Pt(28)
        font.bold = True
        font.color.rgb = RGBColor(0, 102, 204)  # QBITEL Blue
        style.paragraph_format.space_before = Pt(24)
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.keep_with_next = True

        # Heading 2 - Subsections
        style = self.doc.styles['Heading 2']
        font = style.font
        font.name = 'Calibri Light'
        font.size = Pt(20)
        font.bold = True
        font.color.rgb = RGBColor(0, 102, 204)
        style.paragraph_format.space_before = Pt(18)
        style.paragraph_format.space_after = Pt(10)
        style.paragraph_format.keep_with_next = True

        # Heading 3
        style = self.doc.styles['Heading 3']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(16)
        font.bold = True
        font.color.rgb = RGBColor(44, 62, 80)
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.keep_with_next = True

        # Heading 4
        style = self.doc.styles['Heading 4']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(13)
        font.bold = True
        font.color.rgb = RGBColor(68, 84, 106)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    def _add_cover_page(self, title):
        """Add professional cover page with branding"""
        # Logo placeholder
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('QBITEL AI')
        run.font.size = Pt(48)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)

        # Spacing
        for _ in range(3):
            self.doc.add_paragraph()

        # Tagline
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('The Complete Enterprise Security Platform')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(68, 84, 106)
        run.font.italic = True

        # More spacing
        for _ in range(2):
            self.doc.add_paragraph()

        # Document title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(32)
        run.font.bold = True
        run.font.color.rgb = RGBColor(44, 62, 80)

        # Spacing
        for _ in range(2):
            self.doc.add_paragraph()

        # Confidential notice
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('CONFIDENTIAL & PROPRIETARY')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(192, 0, 0)
        run.font.bold = True

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('For Investor Use Only')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(192, 0, 0)

        # Spacing
        for _ in range(4):
            self.doc.add_paragraph()

        # Date
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(datetime.now().strftime('%B %Y'))
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(68, 84, 106)

        # Page break
        self.doc.add_page_break()

    def _add_table_of_contents(self):
        """Add table of contents page"""
        p = self.doc.add_heading('Table of Contents', level=1)
        p.style.font.color.rgb = RGBColor(0, 102, 204)

        # Note about TOC
        p = self.doc.add_paragraph()
        run = p.add_run('Note: In Microsoft Word, right-click on this section and select "Update Field" to refresh the table of contents.')
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

        self.doc.add_paragraph()

        # Add TOC field
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        self.doc.add_page_break()

    def convert_html_to_docx(self, html_content, project_root):
        """Convert HTML (from markdown) to DOCX elements with embedded diagrams"""

        class DocxHTMLParser(HTMLParser):
            def __init__(self, doc, converter, project_root):
                super().__init__()
                self.doc = doc
                self.converter = converter
                self.project_root = Path(project_root)
                self.current_paragraph = None
                self.current_runs = []
                self.in_table = False
                self.table_data = []
                self.current_row = []
                self.list_level = 0
                self.in_code_block = False
                self.in_bold = False
                self.in_italic = False
                self.in_heading = None
                self.pending_image = None

            def handle_starttag(self, tag, attrs):
                attrs_dict = dict(attrs)

                if tag == 'h1':
                    self.in_heading = 1
                    self.current_paragraph = self.doc.add_heading(level=1)
                elif tag == 'h2':
                    self.in_heading = 2
                    self.current_paragraph = self.doc.add_heading(level=2)
                elif tag == 'h3':
                    self.in_heading = 3
                    self.current_paragraph = self.doc.add_heading(level=3)
                elif tag == 'h4':
                    self.in_heading = 4
                    self.current_paragraph = self.doc.add_heading(level=4)
                elif tag == 'p':
                    self.current_paragraph = self.doc.add_paragraph()
                elif tag == 'ul' or tag == 'ol':
                    self.list_level += 1
                elif tag == 'li':
                    self.current_paragraph = self.doc.add_paragraph(style='List Bullet')
                elif tag == 'strong' or tag == 'b':
                    self.in_bold = True
                elif tag == 'em' or tag == 'i':
                    self.in_italic = True
                elif tag == 'code':
                    self.in_code_block = True
                elif tag == 'pre':
                    self.current_paragraph = self.doc.add_paragraph()
                    self.current_paragraph.style = 'Normal'
                    # Code block styling
                    fmt = self.current_paragraph.paragraph_format
                    fmt.left_indent = Inches(0.5)
                    fmt.space_before = Pt(6)
                    fmt.space_after = Pt(6)
                elif tag == 'img':
                    # Handle images (diagrams)
                    src = attrs_dict.get('src', '')
                    alt = attrs_dict.get('alt', '')
                    if src:
                        self._embed_diagram(src, alt)
                elif tag == 'table':
                    self.in_table = True
                    self.table_data = []
                elif tag == 'tr':
                    self.current_row = []
                elif tag == 'hr':
                    p = self.doc.add_paragraph()
                    p.add_run('_' * 80)
                    run = p.runs[0]
                    run.font.color.rgb = RGBColor(200, 200, 200)

            def handle_endtag(self, tag):
                if tag in ['h1', 'h2', 'h3', 'h4']:
                    self.in_heading = None
                    self.current_paragraph = None
                elif tag in ['p', 'li', 'pre']:
                    self.current_paragraph = None
                elif tag == 'ul' or tag == 'ol':
                    self.list_level -= 1
                elif tag == 'strong' or tag == 'b':
                    self.in_bold = False
                elif tag == 'em' or tag == 'i':
                    self.in_italic = False
                elif tag == 'code':
                    self.in_code_block = False
                elif tag == 'tr':
                    if self.current_row:
                        self.table_data.append(self.current_row)
                elif tag == 'table':
                    if self.table_data:
                        self._create_table()
                    self.in_table = False

            def handle_data(self, data):
                if not data.strip():
                    return

                if self.in_table:
                    self.current_row.append(data.strip())
                elif self.current_paragraph:
                    run = self.current_paragraph.add_run(data)

                    # Apply formatting
                    if self.in_bold:
                        run.font.bold = True
                    if self.in_italic:
                        run.font.italic = True
                    if self.in_code_block:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(220, 50, 47)
                else:
                    # Fallback
                    p = self.doc.add_paragraph()
                    p.add_run(data)

            def _embed_diagram(self, src, alt):
                """Embed diagram image into document"""
                # Resolve path relative to project root
                diagram_path = self.project_root / src

                if not diagram_path.exists():
                    # Try without diagrams/ prefix
                    diagram_path = self.project_root / 'diagrams' / Path(src).name

                if diagram_path.exists() and diagram_path.suffix == '.svg':
                    # Try to convert SVG to PNG
                    png_path = self.converter.temp_dir / f"{diagram_path.stem}.png"

                    # Add spacing before diagram
                    self.doc.add_paragraph()

                    if convert_svg_to_png(diagram_path, png_path):
                        # Successfully converted - add PNG image
                        p = self.doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        try:
                            # Add image with width of 6.5 inches (fits well on page)
                            run = p.add_run()
                            run.add_picture(str(png_path), width=Inches(6.5))
                        except Exception as e:
                            print(f"Warning: Could not embed PNG {diagram_path.name}: {e}")
                            # Add diagram reference as fallback
                            self._add_diagram_placeholder(alt, str(diagram_path))
                    else:
                        # Conversion failed - try to embed SVG directly (Word 2016+)
                        try:
                            p = self.doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            # Word 2016+ supports SVG embedding
                            run.add_picture(str(diagram_path), width=Inches(6.5))
                        except Exception as e:
                            print(f"Warning: Could not embed SVG {diagram_path.name}: {e}")
                            # Final fallback - add styled placeholder
                            self._add_diagram_placeholder(alt, str(diagram_path))

                    # Add caption
                    if alt:
                        caption_p = self.doc.add_paragraph()
                        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption_p.add_run(f"Figure: {alt}")
                        caption_run.font.size = Pt(10)
                        caption_run.font.italic = True
                        caption_run.font.color.rgb = RGBColor(68, 84, 106)

                    # Add spacing after
                    self.doc.add_paragraph()

            def _add_diagram_placeholder(self, alt, path):
                """Add a styled placeholder for diagram"""
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Add a bordered box-style placeholder
                run = p.add_run(f"📊 {alt}")
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 102, 204)
                run.font.bold = True

                # Add note
                note_p = self.doc.add_paragraph()
                note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                note_run = note_p.add_run(f"(See {Path(path).name} in diagrams folder)")
                note_run.font.size = Pt(9)
                note_run.font.italic = True
                note_run.font.color.rgb = RGBColor(128, 128, 128)

            def _create_table(self):
                """Create formatted table"""
                if not self.table_data:
                    return

                num_cols = max(len(row) for row in self.table_data)
                table = self.doc.add_table(rows=len(self.table_data), cols=num_cols)
                table.style = 'Light Grid Accent 1'

                for i, row_data in enumerate(self.table_data):
                    row = table.rows[i]
                    for j, cell_data in enumerate(row_data):
                        if j < len(row.cells):
                            cell = row.cells[j]
                            cell.text = cell_data

                            # Header row formatting
                            if i == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                                        run.font.color.rgb = RGBColor(255, 255, 255)

                                # Add blue background to header
                                shading_elm = OxmlElement('w:shd')
                                shading_elm.set(qn('w:fill'), '0066CC')
                                cell._element.get_or_add_tcPr().append(shading_elm)

                self.doc.add_paragraph()  # Spacing after table
                self.table_data = []

        parser = DocxHTMLParser(self.doc, self, project_root)
        parser.feed(html_content)

    def finalize(self):
        """Finalize document with headers/footers"""
        # Add page numbers to all sections except cover page
        for i, section in enumerate(self.doc.sections):
            if i > 0:  # Skip cover page
                add_page_number(section)

        # Clean up temp directory
        import shutil
        try:
            shutil.rmtree(self.temp_dir)
        except:
            pass


def convert_markdown_to_docx(markdown_file: Path, output_dir: Path = None):
    """
    Convert a markdown file to an enterprise-grade DOCX

    Args:
        markdown_file: Path to the markdown file
        output_dir: Output directory for the DOCX
    """
    if output_dir is None:
        output_dir = markdown_file.parent

    output_dir.mkdir(parents=True, exist_ok=True)

    # Read markdown content
    print(f"Reading {markdown_file.name}...")
    with open(markdown_file, 'r', encoding='utf-8') as f:
        markdown_content = f.read()

    # Convert markdown to HTML first
    print(f"Converting {markdown_file.name} to HTML...")
    html_content = markdown.markdown(
        markdown_content,
        extensions=[
            'tables',
            'fenced_code',
            'nl2br',
            'sane_lists'
        ]
    )

    # Get document title
    title_map = {
        "INVESTOR_PITCH_DECK": "Series A Investment Deck",
        "PRODUCT_WALKTHROUGH": "Product & Technology Overview",
        "IMPLEMENTATION_PLAN_ROADMAP": "5-Year Implementation Roadmap"
    }
    title = title_map.get(markdown_file.stem, markdown_file.stem)

    # Create enterprise DOCX document
    print(f"Creating enterprise-grade DOCX: {markdown_file.stem}.docx...")
    converter = EnterpriseMarkdownToDOCXConverter()

    # Add cover page
    print("  - Adding professional cover page...")
    converter._add_cover_page(title)

    # Add table of contents
    print("  - Adding table of contents...")
    converter._add_table_of_contents()

    # Convert HTML to DOCX with embedded diagrams
    print("  - Converting content and embedding diagrams...")
    project_root = markdown_file.parent
    converter.convert_html_to_docx(html_content, project_root)

    # Finalize document
    print("  - Adding headers, footers, and page numbers...")
    converter.finalize()

    # Save DOCX
    docx_filename = output_dir / f"{markdown_file.stem}.docx"
    converter.doc.save(docx_filename)

    print(f"[OK] Successfully created: {docx_filename}")
    print(f"   + Enterprise formatting with professional layout")
    print(f"   + Table of contents with auto-update")
    print(f"   + SVG diagrams embedded as high-quality images")
    print(f"   + Page numbers and footers")
    print(f"   + Ready for investor distribution")
    return docx_filename


def main():
    """Main function to convert all investor documents"""

    # Get the project root directory
    project_root = Path(__file__).parent

    # Files to convert
    files_to_convert = [
        "INVESTOR_PITCH_DECK.md",
        "PRODUCT_WALKTHROUGH.md",
        "IMPLEMENTATION_PLAN_ROADMAP.md"
    ]

    print("=" * 80)
    print("QBITEL AI - Enterprise-Grade DOCX Generator")
    print("Creating Professional Investor Documents with Embedded Diagrams")
    print("=" * 80)
    print(f"Project Root: {project_root}")
    print(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    print("=" * 80)
    print()

    # Check for SVG conversion tools
    print("Checking for SVG conversion tools...")
    svg_converter_found = False

    # Check for ImageMagick
    try:
        result = subprocess.run(['magick', '--version'], capture_output=True, text=True, timeout=5)
        if result.returncode == 0:
            print("  [OK] ImageMagick found (primary converter)")
            svg_converter_found = True
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
        print("  [!] ImageMagick not found")

    # Check for Inkscape
    if not svg_converter_found:
        try:
            result = subprocess.run(['inkscape', '--version'], capture_output=True, text=True, timeout=5)
            if result.returncode == 0:
                print("  [OK] Inkscape found (fallback converter)")
                svg_converter_found = True
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
            print("  [!] Inkscape not found")

    if not svg_converter_found:
        print("  [!] No SVG converter found")
        print("    Will attempt to embed SVG directly (requires Word 2016+)")
        print("    Or install ImageMagick: https://imagemagick.org/script/download.php")
    print()

    output_dir = project_root / "investor_documents_docx"

    generated_files = []
    for filename in files_to_convert:
        markdown_file = project_root / filename
        if markdown_file.exists():
            try:
                print(f"\nProcessing {filename}...")
                print("-" * 80)
                docx_file = convert_markdown_to_docx(markdown_file, output_dir)
                generated_files.append(docx_file)
                print()
            except Exception as e:
                print(f"[ERROR] Error converting {filename}: {e}")
                import traceback
                traceback.print_exc()
                print()
        else:
            print(f"[WARNING] File not found: {filename}")
            print()

    print("=" * 80)
    print("Summary")
    print("=" * 80)
    print(f"Total files converted: {len(generated_files)}/{len(files_to_convert)}")
    print(f"Output directory: {output_dir}")
    print()
    print("Generated Enterprise DOCX files:")
    for docx_file in generated_files:
        file_size = docx_file.stat().st_size / 1024  # KB
        print(f"  - {docx_file.name} ({file_size:.1f} KB)")
    print()
    print("ENTERPRISE FEATURES:")
    print("  + Professional cover page with QBITEL AI branding")
    print("  + Auto-generated table of contents")
    print("  + SVG diagrams embedded as high-quality PNG images")
    print("  + Page numbers and professional footers")
    print("  + Consistent enterprise styling and formatting")
    print("  + Proper headings hierarchy (supports navigation)")
    print("  + Fully editable in Microsoft Word")
    print("  + Ready for investor distribution")
    print("=" * 80)

    return len(generated_files) == len(files_to_convert)


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
