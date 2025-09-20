"""
CRONOS AI - Automated Report Generator

Enterprise-grade compliance report generation with multiple formats,
LLM-powered content generation, and regulator-ready outputs.
"""

import asyncio
import logging
import json
import io
import base64
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, field
from enum import Enum
from datetime import datetime, timedelta
import jinja2
from pathlib import Path

from ..llm.unified_llm_service import UnifiedLLMService, LLMRequest, get_llm_service
from ..core.config import Config
from ..core.exceptions import CronosAIException
from ..monitoring.enterprise_metrics import get_enterprise_metrics
from .regulatory_kb import (
    ComplianceAssessment,
    ComplianceGap, 
    ComplianceRecommendation,
    RequirementSeverity,
    ComplianceFramework
)

logger = logging.getLogger(__name__)

class ReportException(CronosAIException):
    """Report generation specific exception."""
    pass

class ReportFormat(Enum):
    """Supported report formats."""
    PDF = "pdf"
    EXCEL = "excel"
    JSON = "json"
    HTML = "html"
    WORD = "docx"
    CSV = "csv"

class ReportType(Enum):
    """Types of compliance reports."""
    EXECUTIVE_SUMMARY = "executive_summary"
    DETAILED_TECHNICAL = "detailed_technical"
    GAP_ANALYSIS = "gap_analysis"
    REMEDIATION_PLAN = "remediation_plan"
    RISK_ASSESSMENT = "risk_assessment"
    REGULATORY_FILING = "regulatory_filing"

@dataclass
class ReportTemplate:
    """Report template configuration."""
    name: str
    description: str
    format: ReportFormat
    template_path: str
    sections: List[str] = field(default_factory=list)
    required_data: List[str] = field(default_factory=list)
    target_audience: str = "technical"
    regulatory_specific: bool = False
    executive_summary: bool = True
    page_limit: Optional[int] = None

@dataclass
class ComplianceReport:
    """Generated compliance report."""
    report_id: str
    title: str
    framework: str
    report_type: ReportType
    format: ReportFormat
    generated_date: datetime
    content: bytes
    metadata: Dict[str, Any] = field(default_factory=dict)
    file_name: str = ""
    file_size: int = 0
    checksum: str = ""
    audit_trail: Dict[str, Any] = field(default_factory=dict)

@dataclass
class ReportSection:
    """Individual report section."""
    title: str
    content: str
    subsections: List['ReportSection'] = field(default_factory=list)
    charts: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    page_break: bool = False

class ReportTemplateManager:
    """Manages report templates and formatting."""
    
    def __init__(self, config: Config):
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # Template directory
        self.template_dir = Path(__file__).parent / "templates"
        self.template_dir.mkdir(exist_ok=True)
        
        # Jinja2 environment for template rendering
        self.jinja_env = jinja2.Environment(
            loader=jinja2.FileSystemLoader(str(self.template_dir)),
            autoescape=jinja2.select_autoescape(['html', 'xml'])
        )
        
        # Built-in templates
        self.templates = self._load_builtin_templates()
    
    def _load_builtin_templates(self) -> Dict[str, ReportTemplate]:
        """Load built-in report templates."""
        templates = {}
        
        # Executive Summary Template
        templates["executive_summary_pdf"] = ReportTemplate(
            name="Executive Summary (PDF)",
            description="High-level compliance summary for executives",
            format=ReportFormat.PDF,
            template_path="executive_summary.html",
            sections=["overview", "compliance_score", "key_gaps", "recommendations", "next_steps"],
            required_data=["assessment", "framework_info"],
            target_audience="executive",
            executive_summary=True,
            page_limit=5
        )
        
        # Detailed Technical Report
        templates["detailed_technical_pdf"] = ReportTemplate(
            name="Detailed Technical Report (PDF)",
            description="Comprehensive technical compliance analysis",
            format=ReportFormat.PDF,
            template_path="detailed_technical.html",
            sections=["executive_summary", "methodology", "requirements_analysis", "gaps", "recommendations", "appendices"],
            required_data=["assessment", "framework_info", "system_data"],
            target_audience="technical",
            regulatory_specific=True
        )
        
        # Regulatory Filing Template
        templates["regulatory_filing_pdf"] = ReportTemplate(
            name="Regulatory Filing (PDF)",
            description="Formal regulatory submission document",
            format=ReportFormat.PDF,
            template_path="regulatory_filing.html",
            sections=["cover", "attestation", "summary", "detailed_findings", "evidence", "signatures"],
            required_data=["assessment", "framework_info", "organization_info"],
            target_audience="regulator",
            regulatory_specific=True
        )
        
        # Gap Analysis Report
        templates["gap_analysis_excel"] = ReportTemplate(
            name="Gap Analysis (Excel)",
            description="Detailed gap analysis with action items",
            format=ReportFormat.EXCEL,
            template_path="gap_analysis.xlsx",
            sections=["summary", "gaps", "timeline", "resources"],
            required_data=["assessment", "gaps"],
            target_audience="technical"
        )
        
        return templates
    
    def get_template(self, template_name: str) -> Optional[ReportTemplate]:
        """Get specific template by name."""
        return self.templates.get(template_name)
    
    def list_templates(
        self, 
        format_filter: Optional[ReportFormat] = None,
        audience_filter: Optional[str] = None
    ) -> List[ReportTemplate]:
        """List available templates with optional filtering."""
        templates = list(self.templates.values())
        
        if format_filter:
            templates = [t for t in templates if t.format == format_filter]
        
        if audience_filter:
            templates = [t for t in templates if t.target_audience == audience_filter]
        
        return templates

class AutomatedReportGenerator:
    """Main automated report generator with LLM-powered content creation."""
    
    def __init__(
        self, 
        config: Config,
        llm_service: Optional[UnifiedLLMService] = None
    ):
        self.config = config
        self.llm_service = llm_service or get_llm_service()
        self.template_manager = ReportTemplateManager(config)
        self.logger = logging.getLogger(__name__)
        self.metrics = get_enterprise_metrics()
        
        # Report formatters
        self.formatters = {
            ReportFormat.PDF: self._generate_pdf_report,
            ReportFormat.HTML: self._generate_html_report,
            ReportFormat.EXCEL: self._generate_excel_report,
            ReportFormat.JSON: self._generate_json_report,
            ReportFormat.WORD: self._generate_word_report,
            ReportFormat.CSV: self._generate_csv_report
        }
        
        # Initialize formatters
        self._initialize_formatters()
    
    def _initialize_formatters(self):
        """Initialize report formatting libraries."""
        try:
            # Try to import optional dependencies
            global weasyprint, openpyxl, python_docx, pandas
            import weasyprint
            import openpyxl
            import docx as python_docx
            import pandas
            self.logger.info("All report formatters initialized successfully")
        except ImportError as e:
            self.logger.warning(f"Some report formatters unavailable: {e}")
    
    async def generate_compliance_report(
        self, 
        assessment: ComplianceAssessment,
        report_type: ReportType = ReportType.DETAILED_TECHNICAL,
        format: ReportFormat = ReportFormat.PDF,
        template_name: Optional[str] = None,
        custom_sections: Optional[List[str]] = None
    ) -> ComplianceReport:
        """
        Generate comprehensive compliance report.
        
        Args:
            assessment: Compliance assessment data
            report_type: Type of report to generate
            format: Output format
            template_name: Specific template to use
            custom_sections: Custom sections to include
            
        Returns:
            Generated compliance report
        """
        try:
            start_time = datetime.utcnow()
            self.logger.info(f"Generating {report_type.value} report in {format.value} format")
            
            # Select appropriate template
            template = self._select_template(template_name, report_type, format)
            
            # Generate LLM-powered content
            report_content = await self._generate_report_content(
                assessment, template, custom_sections
            )
            
            # Format report using appropriate formatter
            formatted_report = await self._format_report(
                report_content, format, template
            )
            
            # Create report metadata
            report_id = self._generate_report_id(assessment.framework, report_type)
            title = self._generate_report_title(assessment, report_type)
            
            # Create compliance report object
            compliance_report = ComplianceReport(
                report_id=report_id,
                title=title,
                framework=assessment.framework,
                report_type=report_type,
                format=format,
                generated_date=start_time,
                content=formatted_report,
                file_name=f"{report_id}.{format.value}",
                file_size=len(formatted_report),
                metadata={
                    'template': template.name if template else 'custom',
                    'assessment_date': assessment.assessment_date.isoformat(),
                    'compliance_score': assessment.overall_compliance_score,
                    'risk_score': assessment.risk_score,
                    'generator_version': '1.0.0'
                },
                audit_trail={
                    'generated_by': 'automated_system',
                    'generation_time': (datetime.utcnow() - start_time).total_seconds(),
                    'llm_provider': 'unified_llm_service',
                    'template_used': template.name if template else 'custom'
                }
            )
            
            # Record metrics
            self.metrics.record_protocol_discovery_metric(
                "compliance_report_generation_duration_seconds",
                (datetime.utcnow() - start_time).total_seconds(),
                {
                    "framework": assessment.framework,
                    "report_type": report_type.value,
                    "format": format.value
                }
            )
            
            self.metrics.increment_protocol_discovery_counter(
                "compliance_reports_generated_total",
                labels={
                    "framework": assessment.framework,
                    "report_type": report_type.value,
                    "format": format.value
                }
            )
            
            self.logger.info(f"Report generated successfully: {report_id}")
            return compliance_report
            
        except Exception as e:
            self.logger.error(f"Report generation failed: {e}")
            self.metrics.increment_protocol_discovery_counter(
                "compliance_report_errors_total",
                labels={
                    "framework": assessment.framework,
                    "error_type": type(e).__name__
                }
            )
            raise ReportException(f"Report generation failed: {e}")
    
    def _select_template(
        self, 
        template_name: Optional[str],
        report_type: ReportType,
        format: ReportFormat
    ) -> Optional[ReportTemplate]:
        """Select appropriate template for report generation."""
        if template_name:
            return self.template_manager.get_template(template_name)
        
        # Auto-select template based on report type and format
        template_mapping = {
            (ReportType.EXECUTIVE_SUMMARY, ReportFormat.PDF): "executive_summary_pdf",
            (ReportType.DETAILED_TECHNICAL, ReportFormat.PDF): "detailed_technical_pdf",
            (ReportType.REGULATORY_FILING, ReportFormat.PDF): "regulatory_filing_pdf",
            (ReportType.GAP_ANALYSIS, ReportFormat.EXCEL): "gap_analysis_excel"
        }
        
        template_key = (report_type, format)
        template_name = template_mapping.get(template_key)
        
        return self.template_manager.get_template(template_name) if template_name else None
    
    async def _generate_report_content(
        self,
        assessment: ComplianceAssessment,
        template: Optional[ReportTemplate],
        custom_sections: Optional[List[str]]
    ) -> Dict[str, Any]:
        """Generate report content using LLM."""
        try:
            # Determine sections to include
            sections = custom_sections or (template.sections if template else [
                "executive_summary", "compliance_overview", "gap_analysis", 
                "recommendations", "next_steps"
            ])
            
            # Generate content for each section
            report_content = {
                'metadata': {
                    'framework': assessment.framework,
                    'assessment_date': assessment.assessment_date.isoformat(),
                    'generation_date': datetime.utcnow().isoformat(),
                    'compliance_score': assessment.overall_compliance_score,
                    'risk_score': assessment.risk_score
                },
                'sections': {}
            }
            
            # Generate sections concurrently for performance
            section_tasks = [
                self._generate_section_content(section, assessment)
                for section in sections
            ]
            
            section_results = await asyncio.gather(*section_tasks, return_exceptions=True)
            
            # Collect successful results
            for section, result in zip(sections, section_results):
                if isinstance(result, Exception):
                    self.logger.error(f"Failed to generate section {section}: {result}")
                    report_content['sections'][section] = {
                        'title': section.replace('_', ' ').title(),
                        'content': f"Error generating section: {str(result)}"
                    }
                else:
                    report_content['sections'][section] = result
            
            return report_content
            
        except Exception as e:
            self.logger.error(f"Failed to generate report content: {e}")
            raise ReportException(f"Content generation failed: {e}")
    
    async def _generate_section_content(
        self,
        section_name: str,
        assessment: ComplianceAssessment
    ) -> Dict[str, Any]:
        """Generate content for a specific report section using LLM."""
        try:
            # Create section-specific prompt
            prompt = self._create_section_prompt(section_name, assessment)
            
            # Request LLM content generation
            llm_request = LLMRequest(
                prompt=prompt,
                feature_domain="compliance_reporter",
                context={
                    'section': section_name,
                    'framework': assessment.framework,
                    'compliance_score': assessment.overall_compliance_score
                },
                max_tokens=2000,
                temperature=0.2  # Low temperature for consistent professional content
            )
            
            response = await self.llm_service.process_request(llm_request)
            
            # Parse response and structure content
            section_content = self._parse_section_response(section_name, response.content)
            
            return section_content
            
        except Exception as e:
            self.logger.error(f"Failed to generate section {section_name}: {e}")
            return {
                'title': section_name.replace('_', ' ').title(),
                'content': f"Section generation failed: {str(e)}",
                'error': True
            }
    
    def _create_section_prompt(
        self,
        section_name: str,
        assessment: ComplianceAssessment
    ) -> str:
        """Create LLM prompt for specific report section."""
        base_context = f"""
Framework: {assessment.framework}
Assessment Date: {assessment.assessment_date.strftime('%Y-%m-%d')}
Overall Compliance Score: {assessment.overall_compliance_score:.1f}%
Risk Score: {assessment.risk_score:.1f}%
Total Requirements: {assessment.compliant_requirements + assessment.non_compliant_requirements + assessment.partially_compliant_requirements}
Compliant: {assessment.compliant_requirements}
Non-Compliant: {assessment.non_compliant_requirements}
Partially Compliant: {assessment.partially_compliant_requirements}
Critical Gaps: {len([g for g in assessment.gaps if g.severity == RequirementSeverity.CRITICAL])}
"""
        
        section_prompts = {
            'executive_summary': f"""
Write a professional executive summary for a {assessment.framework} compliance assessment report.

{base_context}

Create a concise, executive-level summary that includes:
1. Overall compliance status and key findings
2. Critical compliance gaps requiring immediate attention
3. Business risk implications
4. High-level recommendations and next steps
5. Timeline for achieving full compliance

Target audience: C-level executives and board members
Style: Professional, concise, business-focused
Length: 2-3 paragraphs maximum
""",
            
            'compliance_overview': f"""
Write a comprehensive compliance overview section for a {assessment.framework} assessment report.

{base_context}

Include:
1. Methodology used for the assessment
2. Scope and limitations
3. Overall compliance posture
4. Key compliance metrics and trends
5. Comparison to industry benchmarks (if applicable)

Style: Professional, technical but accessible
""",
            
            'gap_analysis': f"""
Write a detailed gap analysis section for a {assessment.framework} compliance report.

{base_context}

Top Critical Gaps:
{chr(10).join([f"- {gap.requirement_title}: {gap.gap_description}" for gap in assessment.gaps[:5] if gap.severity == RequirementSeverity.CRITICAL])}

Include:
1. Summary of identified gaps by severity
2. Root cause analysis of major gaps
3. Impact assessment for each critical gap
4. Dependencies between gaps
5. Risk implications of unaddressed gaps

Style: Analytical, detailed, risk-focused
""",
            
            'recommendations': f"""
Write a recommendations section for a {assessment.framework} compliance report.

{base_context}

Priority Recommendations:
{chr(10).join([f"- {rec.title}: {rec.description}" for rec in assessment.recommendations[:5]])}

Include:
1. Prioritized list of recommendations
2. Implementation approach for each recommendation
3. Resource requirements and estimated costs
4. Expected timeline for implementation
5. Success metrics and validation criteria

Style: Actionable, practical, solution-oriented
""",
            
            'next_steps': f"""
Write a next steps section for a {assessment.framework} compliance report.

{base_context}

Next Assessment Due: {assessment.next_assessment_due.strftime('%Y-%m-%d')}

Include:
1. Immediate actions required (next 30 days)
2. Short-term initiatives (next 90 days)
3. Long-term compliance roadmap (6-12 months)
4. Governance and oversight structure
5. Continuous monitoring and improvement process

Style: Action-oriented, specific, time-bound
"""
        }
        
        return section_prompts.get(section_name, f"""
Write a professional section about {section_name} for a {assessment.framework} compliance report.
{base_context}
Provide comprehensive, accurate, and actionable content appropriate for enterprise compliance reporting.
""")
    
    def _parse_section_response(self, section_name: str, content: str) -> Dict[str, Any]:
        """Parse LLM response into structured section content."""
        return {
            'title': section_name.replace('_', ' ').title(),
            'content': content.strip(),
            'subsections': self._extract_subsections(content),
            'generated_at': datetime.utcnow().isoformat()
        }
    
    def _extract_subsections(self, content: str) -> List[Dict[str, str]]:
        """Extract subsections from generated content."""
        subsections = []
        lines = content.split('\n')
        current_subsection = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if line.startswith('#') or line.startswith('**') and line.endswith('**'):
                # New subsection found
                if current_subsection:
                    subsections.append({
                        'title': current_subsection,
                        'content': '\n'.join(current_content).strip()
                    })
                current_subsection = line.replace('#', '').replace('**', '').strip()
                current_content = []
            else:
                current_content.append(line)
        
        # Add final subsection
        if current_subsection:
            subsections.append({
                'title': current_subsection,
                'content': '\n'.join(current_content).strip()
            })
        
        return subsections
    
    async def _format_report(
        self,
        content: Dict[str, Any],
        format: ReportFormat,
        template: Optional[ReportTemplate]
    ) -> bytes:
        """Format report content into specified output format."""
        formatter = self.formatters.get(format)
        if not formatter:
            raise ReportException(f"Unsupported format: {format.value}")
        
        return await formatter(content, template)
    
    async def _generate_pdf_report(
        self,
        content: Dict[str, Any],
        template: Optional[ReportTemplate]
    ) -> bytes:
        """Generate PDF report using HTML template and WeasyPrint."""
        try:
            # Generate HTML first
            html_content = await self._generate_html_content(content, template)
            
            # Convert HTML to PDF using WeasyPrint
            import weasyprint
            
            # Create PDF
            pdf_document = weasyprint.HTML(string=html_content)
            pdf_bytes = pdf_document.write_pdf()
            
            return pdf_bytes
            
        except ImportError:
            raise ReportException("WeasyPrint not available for PDF generation")
        except Exception as e:
            raise ReportException(f"PDF generation failed: {e}")
    
    async def _generate_html_report(
        self,
        content: Dict[str, Any],
        template: Optional[ReportTemplate]
    ) -> bytes:
        """Generate HTML report."""
        try:
            html_content = await self._generate_html_content(content, template)
            return html_content.encode('utf-8')
        except Exception as e:
            raise ReportException(f"HTML generation failed: {e}")
    
    async def _generate_html_content(
        self,
        content: Dict[str, Any],
        template: Optional[ReportTemplate]
    ) -> str:
        """Generate HTML content from report data."""
        if template and template.template_path:
            # Use custom template
            try:
                jinja_template = self.template_manager.jinja_env.get_template(template.template_path)
                return jinja_template.render(content=content)
            except jinja2.TemplateNotFound:
                self.logger.warning(f"Template not found: {template.template_path}, using default")
        
        # Use default HTML template
        html_template = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ content.metadata.framework }} Compliance Report</title>
    <style>
        body { font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }
        .header { border-bottom: 2px solid #333; padding-bottom: 20px; margin-bottom: 30px; }
        .metadata { background: #f5f5f5; padding: 15px; margin-bottom: 20px; }
        .section { margin-bottom: 30px; }
        .section-title { color: #333; border-bottom: 1px solid #ccc; padding-bottom: 5px; }
        .subsection { margin-left: 20px; margin-bottom: 15px; }
        .score { font-size: 24px; font-weight: bold; color: #0066cc; }
        .risk-high { color: #dc3545; }
        .risk-medium { color: #ffc107; }
        .risk-low { color: #28a745; }
        table { width: 100%; border-collapse: collapse; margin: 15px 0; }
        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
        th { background-color: #f2f2f2; }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ content.metadata.framework }} Compliance Assessment Report</h1>
        <div class="metadata">
            <p><strong>Assessment Date:</strong> {{ content.metadata.assessment_date }}</p>
            <p><strong>Report Generated:</strong> {{ content.metadata.generation_date }}</p>
            <p><strong>Overall Compliance Score:</strong> <span class="score">{{ "%.1f"|format(content.metadata.compliance_score) }}%</span></p>
            <p><strong>Risk Score:</strong> <span class="score risk-{{ 'high' if content.metadata.risk_score > 70 else 'medium' if content.metadata.risk_score > 30 else 'low' }}">{{ "%.1f"|format(content.metadata.risk_score) }}%</span></p>
        </div>
    </div>
    
    {% for section_name, section_data in content.sections.items() %}
    <div class="section">
        <h2 class="section-title">{{ section_data.title }}</h2>
        <div class="content">
            {{ section_data.content | replace('\n', '<br>') | safe }}
        </div>
        
        {% if section_data.subsections %}
        {% for subsection in section_data.subsections %}
        <div class="subsection">
            <h3>{{ subsection.title }}</h3>
            <p>{{ subsection.content | replace('\n', '<br>') | safe }}</p>
        </div>
        {% endfor %}
        {% endif %}
    </div>
    {% endfor %}
    
    <div class="footer">
        <p><em>This report was generated automatically by CRONOS AI Compliance Reporter on {{ content.metadata.generation_date }}.</em></p>
    </div>
</body>
</html>
        """
        
        jinja_template = jinja2.Template(html_template)
        return jinja_template.render(content=content)
    
    async def _generate_excel_report(
        self,
        content: Dict[str, Any],
        template: Optional[ReportTemplate]
    ) -> bytes:
        """Generate Excel report."""
        try:
            import openpyxl
            import io
            
            # Create workbook
            wb = openpyxl.Workbook()
            
            # Remove default sheet
            wb.remove(wb.active)
            
            # Create summary sheet
            summary_ws = wb.create_sheet("Summary")
            self._populate_excel_summary(summary_ws, content)
            
            # Create detailed sheets for each section
            for section_name, section_data in content['sections'].items():
                ws = wb.create_sheet(section_name.replace('_', ' ').title()[:31])  # Excel sheet name limit
                self._populate_excel_section(ws, section_data)
            
            # Save to bytes
            excel_buffer = io.BytesIO()
            wb.save(excel_buffer)
            excel_buffer.seek(0)
            
            return excel_buffer.read()
            
        except ImportError:
            raise ReportException("openpyxl not available for Excel generation")
        except Exception as e:
            raise ReportException(f"Excel generation failed: {e}")
    
    def _populate_excel_summary(self, worksheet, content: Dict[str, Any]):
        """Populate Excel summary sheet."""
        metadata = content['metadata']
        
        # Headers
        worksheet['A1'] = 'Compliance Assessment Summary'
        worksheet['A1'].font = openpyxl.styles.Font(size=16, bold=True)
        
        # Metadata
        row = 3
        for key, value in metadata.items():
            worksheet[f'A{row}'] = key.replace('_', ' ').title()
            worksheet[f'B{row}'] = str(value)
            row += 1
    
    def _populate_excel_section(self, worksheet, section_data: Dict[str, Any]):
        """Populate Excel section sheet."""
        worksheet['A1'] = section_data['title']
        worksheet['A1'].font = openpyxl.styles.Font(size=14, bold=True)
        
        # Content
        lines = section_data['content'].split('\n')
        for i, line in enumerate(lines):
            worksheet[f'A{i+3}'] = line
    
    async def _generate_json_report(
        self,
        content: Dict[str, Any],
        template: Optional[ReportTemplate]
    ) -> bytes:
        """Generate JSON report."""
        try:
            json_content = json.dumps(content, indent=2, default=str)
            return json_content.encode('utf-8')
        except Exception as e:
            raise ReportException(f"JSON generation failed: {e}")
    
    async def _generate_word_report(
        self,
        content: Dict[str, Any],
        template: Optional[ReportTemplate]
    ) -> bytes:
        """Generate Word document report."""
        try:
            import docx
            import io
            
            # Create document
            doc = docx.Document()
            
            # Add title
            title = doc.add_heading(f"{content['metadata']['framework']} Compliance Report", 0)
            
            # Add metadata
            metadata_para = doc.add_paragraph()
            metadata = content['metadata']
            for key, value in metadata.items():
                metadata_para.add_run(f"{key.replace('_', ' ').title()}: {value}\n")
            
            # Add sections
            for section_name, section_data in content['sections'].items():
                doc.add_heading(section_data['title'], level=1)
                doc.add_paragraph(section_data['content'])
                
                # Add subsections
                for subsection in section_data.get('subsections', []):
                    doc.add_heading(subsection['title'], level=2)
                    doc.add_paragraph(subsection['content'])
            
            # Save to bytes
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer.read()
            
        except ImportError:
            raise ReportException("python-docx not available for Word generation")
        except Exception as e:
            raise ReportException(f"Word generation failed: {e}")
    
    async def _generate_csv_report(
        self,
        content: Dict[str, Any],
        template: Optional[ReportTemplate]
    ) -> bytes:
        """Generate CSV report."""
        try:
            import io
            import csv
            
            csv_buffer = io.StringIO()
            writer = csv.writer(csv_buffer)
            
            # Write metadata
            writer.writerow(['Metadata'])
            metadata = content['metadata']
            for key, value in metadata.items():
                writer.writerow([key.replace('_', ' ').title(), str(value)])
            
            writer.writerow([])  # Empty row
            
            # Write sections
            for section_name, section_data in content['sections'].items():
                writer.writerow([section_data['title']])
                # Split content into lines and write each as a row
                lines = section_data['content'].split('\n')
                for line in lines:
                    if line.strip():
                        writer.writerow([line.strip()])
                writer.writerow([])  # Empty row between sections
            
            return csv_buffer.getvalue().encode('utf-8')
            
        except Exception as e:
            raise ReportException(f"CSV generation failed: {e}")
    
    def _generate_report_id(self, framework: str, report_type: ReportType) -> str:
        """Generate unique report ID."""
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        return f"{framework}_{report_type.value}_{timestamp}"
    
    def _generate_report_title(self, assessment: ComplianceAssessment, report_type: ReportType) -> str:
        """Generate report title."""
        type_titles = {
            ReportType.EXECUTIVE_SUMMARY: "Executive Summary",
            ReportType.DETAILED_TECHNICAL: "Detailed Technical Assessment",
            ReportType.GAP_ANALYSIS: "Compliance Gap Analysis",
            ReportType.REMEDIATION_PLAN: "Remediation Plan",
            ReportType.RISK_ASSESSMENT: "Risk Assessment",
            ReportType.REGULATORY_FILING: "Regulatory Filing"
        }
        
        type_title = type_titles.get(report_type, "Compliance Report")
        return f"{assessment.framework} {type_title} - {assessment.assessment_date.strftime('%Y-%m-%d')}"
    
    async def generate_multiple_formats(
        self,
        assessment: ComplianceAssessment,
        formats: List[ReportFormat],
        report_type: ReportType = ReportType.DETAILED_TECHNICAL
    ) -> List[ComplianceReport]:
        """Generate reports in multiple formats simultaneously."""
        try:
            # Generate reports for all formats concurrently
            report_tasks = [
                self.generate_compliance_report(assessment, report_type, format)
                for format in formats
            ]
            
            reports = await asyncio.gather(*report_tasks, return_exceptions=True)
            
            # Filter successful reports
            successful_reports = []
            for report in reports:
                if isinstance(report, Exception):
                    self.logger.error(f"Multi-format report generation failed: {report}")
                else:
                    successful_reports.append(report)
            
            return successful_reports
            
        except Exception as e:
            self.logger.error(f"Multi-format report generation failed: {e}")
            raise ReportException(f"Multi-format generation failed: {e}")