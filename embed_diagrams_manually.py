#!/usr/bin/env python3
"""
Manual diagram embedder - opens SVG files in browser and allows screenshot/copy
Then embeds them into DOCX files
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
import webbrowser
import time

def open_svg_in_browser(svg_path):
    """Open SVG file in default browser"""
    webbrowser.open(f'file:///{svg_path.absolute()}')

def embed_image_in_docx(docx_path, image_path, search_text):
    """Find placeholder text and replace with image"""
    doc = Document(docx_path)

    for para in doc.paragraphs:
        if search_text in para.text:
            # Clear the paragraph
            para.clear()
            # Add the image
            run = para.add_run()
            run.add_picture(str(image_path), width=Inches(6.5))
            para.alignment = 1  # Center
            break

    doc.save(docx_path)
    print(f"Updated {docx_path.name}")

def main():
    """Guide user through manual embedding process"""
    project_root = Path(__file__).parent
    diagrams_dir = project_root / 'diagrams'
    docx_dir = project_root / 'investor_documents_docx'

    # SVG files and their corresponding DOCX files
    diagram_mapping = {
        '01_system_architecture.svg': ['INVESTOR_PITCH_DECK.docx'],
        '02_ai_agent_ecosystem.svg': ['INVESTOR_PITCH_DECK.docx'],
        '05_translation_studio_workflow.svg': ['PRODUCT_WALKTHROUGH.docx'],
    }

    print("=" * 80)
    print("Manual Diagram Embedding Guide")
    print("=" * 80)
    print()
    print("Since we don't have ImageMagick installed, here's how to embed diagrams:")
    print()
    print("METHOD 1: Open DOCX, Insert SVG directly (Word 2016+)")
    print("-" * 80)
    print("1. Open the DOCX file in Microsoft Word")
    print("2. Find the diagram placeholder (e.g., '📊 QBITEL AI Platform Architecture')")
    print("3. Click Insert > Pictures > This Device")
    print("4. Navigate to the 'diagrams' folder")
    print("5. Select the SVG file and click Insert")
    print("6. Resize if needed (recommended: 6.5 inches width)")
    print("7. Delete the placeholder text")
    print("8. Save the document")
    print()
    print("METHOD 2: Convert SVG to PNG online")
    print("-" * 80)
    print("1. Go to https://cloudconvert.com/svg-to-png")
    print("2. Upload the SVG file from the 'diagrams' folder")
    print("3. Download the converted PNG")
    print("4. Insert the PNG into Word (same steps as above)")
    print()
    print("FILES TO UPDATE:")
    print("-" * 80)
    for svg_file, docx_files in diagram_mapping.items():
        print(f"\n{svg_file}:")
        print(f"  Location: diagrams\\{svg_file}")
        print(f"  Insert into:")
        for docx in docx_files:
            print(f"    - investor_documents_docx\\{docx}")
    print()
    print("=" * 80)

    # Ask if user wants to open files
    response = input("\nWould you like to open the diagram files in your browser? (y/n): ")
    if response.lower() == 'y':
        for svg_file in diagram_mapping.keys():
            svg_path = diagrams_dir / svg_file
            if svg_path.exists():
                print(f"Opening {svg_file}...")
                open_svg_in_browser(svg_path)
                time.sleep(1)

        print("\nDiagrams opened in browser. You can:")
        print("1. Right-click and 'Save Image As' to save as PNG")
        print("2. Take a screenshot")
        print("3. Or use the SVG files directly in Word 2016+")

if __name__ == "__main__":
    main()
